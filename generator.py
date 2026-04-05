from docx import Document
from datetime import datetime
from google import genai
from dotenv import load_dotenv
import os

load_dotenv()


def generate_architecture(client_info):
    client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))
    prompt = f"""
You are a Technical Sales Engineer.
Based on this client info, propose a specific technical integration architecture in 3-4 sentences.
Mention their actual tools. Be specific, no bullet points, plain text only.

Pain Points: {client_info['pain_points']}
Current Tech Stack: {client_info['current_tech_stack']}
Desired Outcome: {client_info['desired_outcome']}
"""
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt
    )
    return response.text.strip()


def generate_solution_document(client_info, similar_cases):

    print("Generating Solution Design Document...")

    os.makedirs("outputs", exist_ok=True)

    doc = Document()

    # Title
    doc.add_heading(
        "Enterprise Integration Solution Design Document",
        level=1
    )
    doc.add_paragraph(f"Generated on: {datetime.now()}")

    # Section 1 — Pain Points
    doc.add_heading("1. Client Pain Points", level=2)
    for p in client_info["pain_points"]:
        doc.add_paragraph(p, style="List Bullet")

    # Section 2 — Current Tech Stack
    doc.add_heading("2. Current Tech Stack", level=2)
    for t in client_info["current_tech_stack"]:
        doc.add_paragraph(t, style="List Bullet")

    # Section 3 — Desired Outcome
    doc.add_heading("3. Desired Outcome", level=2)
    doc.add_paragraph(client_info["desired_outcome"])

    # Section 4 — Similar Past Projects
    doc.add_heading("4. Similar Past Projects", level=2)
    for case in similar_cases:
        doc.add_paragraph(case[:400])

    # Section 5 — Proposed Architecture
    doc.add_heading("5. Proposed Architecture", level=2)
    architecture_text = generate_architecture(client_info)
    doc.add_paragraph(architecture_text)

    # Save file
    output_file = "outputs/solution_design_document.docx"
    doc.save(output_file)
    print(f"✅ Document saved: {output_file}")
    return output_file